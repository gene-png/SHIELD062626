"""Smoke tests for the CSF PDF + XLSX exporters."""

from __future__ import annotations

import io
import uuid

import pytest

from app.csf.catalog import SUBCATEGORIES
from app.csf.exporters import build_context, render_docx, render_pdf, render_xlsx
from app.csf.gap import analyze as analyze_gaps
from app.csf.maturity import TIER_DEFINITIONS
from app.csf.scoring import compute as compute_score
from app.models.csf_assessment import CsfAnswer, CsfAssessment, CsfAssessmentStatus
from app.models.csf_profile import CsfGapAction


def _build_inputs(*, answers_tier: int | None = 3) -> tuple[CsfAssessment, list[CsfAnswer]]:
    a = CsfAssessment(
        id=uuid.uuid4(),
        service_id=uuid.uuid4(),
        version=1,
        status=CsfAssessmentStatus.APPROVED,
    )
    answers = []
    for sc in SUBCATEGORIES:
        ans = CsfAnswer(
            id=uuid.uuid4(),
            assessment_id=a.id,
            subcategory_code=sc.code,
            maturity_tier=answers_tier,
        )
        answers.append(ans)
    return a, answers


@pytest.fixture()
def context_with_full_tier3():
    a, answers = _build_inputs(answers_tier=3)
    tier_map = {ans.subcategory_code: ans.maturity_tier for ans in answers}
    score = compute_score(tier_map)
    gap = analyze_gaps(tier_map, target_tier=4)  # everything is below T4
    return build_context(
        client_legal_name="Atlas Defense Solutions",
        service_title="NIST CSF 2.0 Assessment",
        assessment=a,
        answers=answers,
        score=score,
        gap=gap,
    )


@pytest.mark.unit
def test_xlsx_render_has_three_sheets(context_with_full_tier3) -> None:
    from openpyxl import load_workbook

    raw = render_xlsx(context_with_full_tier3)
    assert isinstance(raw, bytes)
    assert raw[:2] == b"PK"  # XLSX is a zip envelope
    wb = load_workbook(io.BytesIO(raw))
    assert set(wb.sheetnames) == {"Score Summary", "Answers", "Gap Plan"}


@pytest.mark.unit
def test_xlsx_score_summary_carries_overall_label(context_with_full_tier3) -> None:
    from openpyxl import load_workbook

    raw = render_xlsx(context_with_full_tier3)
    wb = load_workbook(io.BytesIO(raw))
    ws = wb["Score Summary"]
    cells = [(ws.cell(row=r, column=1).value, ws.cell(row=r, column=2).value) for r in range(1, 7)]
    label = dict(cells).get("Overall maturity")
    # All tier 3 -> Repeatable.
    assert label == "Repeatable"


@pytest.mark.unit
def test_xlsx_answers_sheet_has_one_row_per_subcategory(context_with_full_tier3) -> None:
    from openpyxl import load_workbook

    raw = render_xlsx(context_with_full_tier3)
    wb = load_workbook(io.BytesIO(raw))
    ws = wb["Answers"]
    # 106 subcategories + 1 header = 107 rows.
    assert ws.max_row == 107


@pytest.mark.unit
def test_xlsx_gap_plan_contains_top_gaps_when_target_is_adaptive(
    context_with_full_tier3,
) -> None:
    from openpyxl import load_workbook

    raw = render_xlsx(context_with_full_tier3)
    wb = load_workbook(io.BytesIO(raw))
    ws = wb["Gap Plan"]
    # Default top_n=20 with target=4 produces 20 gap rows + 1 header.
    assert ws.max_row == 21


@pytest.mark.unit
def test_pdf_render_produces_valid_pdf(context_with_full_tier3) -> None:
    raw = render_pdf(context_with_full_tier3)
    assert isinstance(raw, bytes)
    assert raw.startswith(b"%PDF-")
    assert len(raw) > 2000


def _pdf_text(raw: bytes) -> str:
    from pypdf import PdfReader

    reader = PdfReader(io.BytesIO(raw))
    return "".join(page.extract_text() for page in reader.pages)


@pytest.mark.unit
def test_pdf_render_carries_title_client_and_a_known_value(context_with_full_tier3) -> None:
    # SMOKE §10: upgrade from %PDF- magic to real content — title, client, and
    # one known computed value (all-tier-3 rolls up to "Repeatable").
    text = _pdf_text(render_pdf(context_with_full_tier3))
    assert "NIST CSF 2.0 Assessment" in text  # service title
    assert "Atlas Defense Solutions" in text  # client name
    assert "Repeatable" in text  # overall maturity label (tier 3)


@pytest.mark.unit
def test_pdf_handles_empty_gap_list() -> None:
    a, answers = _build_inputs(answers_tier=4)  # everyone Adaptive
    tier_map = {ans.subcategory_code: ans.maturity_tier for ans in answers}
    score = compute_score(tier_map)
    # Target tier 3 means everyone is past it -> zero gaps.
    gap = analyze_gaps(tier_map, target_tier=3)
    assert gap.total_gap_count == 0
    ctx = build_context(
        client_legal_name=None,
        service_title="x",
        assessment=a,
        answers=answers,
        score=score,
        gap=gap,
    )
    raw = render_pdf(ctx)
    assert raw.startswith(b"%PDF-")


@pytest.mark.unit
def test_xlsx_handles_empty_gap_list_with_placeholder_row() -> None:
    from openpyxl import load_workbook

    a, answers = _build_inputs(answers_tier=4)
    tier_map = {ans.subcategory_code: ans.maturity_tier for ans in answers}
    score = compute_score(tier_map)
    gap = analyze_gaps(tier_map, target_tier=3)
    ctx = build_context(
        client_legal_name=None,
        service_title="x",
        assessment=a,
        answers=answers,
        score=score,
        gap=gap,
    )
    wb = load_workbook(io.BytesIO(render_xlsx(ctx)))
    ws = wb["Gap Plan"]
    # Header + 1 placeholder row.
    assert ws.max_row == 2
    assert ws.cell(row=2, column=4).value == "No gaps at target tier"


@pytest.mark.unit
def test_build_context_falls_back_to_client_when_legal_name_is_none() -> None:
    a, answers = _build_inputs()
    tier_map = {ans.subcategory_code: ans.maturity_tier for ans in answers}
    score = compute_score(tier_map)
    gap = analyze_gaps(tier_map)
    ctx = build_context(
        client_legal_name=None,
        service_title="x",
        assessment=a,
        answers=answers,
        score=score,
        gap=gap,
    )
    assert ctx.client_legal_name == "Client"


@pytest.mark.unit
def test_pdf_escapes_ampersand_in_header() -> None:
    # The header lines feed reportlab Paragraph markup: a bare "&" in the
    # client name must be escaped so it renders literally instead of
    # re-emitting as an unknown entity with a synthesized semicolon.
    a, answers = _build_inputs(answers_tier=3)
    tier_map = {ans.subcategory_code: ans.maturity_tier for ans in answers}
    score = compute_score(tier_map)
    gap = analyze_gaps(tier_map, target_tier=4)
    ctx = build_context(
        client_legal_name="Rook&Pawn Security",
        service_title="NIST CSF 2.0",
        assessment=a,
        answers=answers,
        score=score,
        gap=gap,
    )
    text = _pdf_text(render_pdf(ctx))
    assert "Rook&Pawn Security" in text


# ---------------------------------------------------------------------------
# S3: POA&M action plan, tier-model methodology, heatmap fills, evidence
# ---------------------------------------------------------------------------


def _pdf_norm(raw: bytes) -> str:
    """Extracted PDF text with reportlab's line wraps collapsed to spaces."""
    return " ".join(_pdf_text(raw).split())


def _docx_norm(raw: bytes) -> str:
    """Every DOCX paragraph AND table cell, wraps collapsed. Tables are not in
    `Document.paragraphs`, so a paragraph-only reader cannot see the Action
    Plan rows at all."""
    from docx import Document

    doc = Document(io.BytesIO(raw))
    chunks = [p.text for p in doc.paragraphs]
    for table in doc.tables:
        for row in table.rows:
            chunks.extend(cell.text for cell in row.cells)
    return " ".join(" ".join(chunks).split())


def _action(code: str, **fields) -> CsfGapAction:
    return CsfGapAction(id=uuid.uuid4(), subcategory_code=code, **fields)


def _ctx_with_actions(actions: dict[str, CsfGapAction] | None = None, **kw):
    """All-tier-3 answers against a T4 target, so every subcategory is a gap."""
    a, answers = _build_inputs(answers_tier=3)
    tier_map = {ans.subcategory_code: ans.maturity_tier for ans in answers}
    return build_context(
        client_legal_name="Atlas Defense Solutions",
        service_title="NIST CSF 2.0 Assessment",
        assessment=a,
        answers=answers,
        score=compute_score(tier_map),
        gap=analyze_gaps(tier_map, target_tier=4),
        actions=actions or {},
        **kw,
    )


GAP_PLAN_HEADERS = [
    "Subcategory",
    "Function",
    "Category",
    "Name",
    "Current tier",
    "Target tier",
    "Gap size",
    "Priority",
    "Characterization",
    "Owner",
    "Deadline",
    "Resources",
    "Success criteria",
    "POA&M ref",
    "Notes",
]


@pytest.mark.unit
def test_xlsx_gap_plan_header_contract() -> None:
    from openpyxl import load_workbook

    ws = load_workbook(io.BytesIO(render_xlsx(_ctx_with_actions())))["Gap Plan"]
    got = [ws.cell(row=1, column=c).value for c in range(1, ws.max_column + 1)]
    assert got == GAP_PLAN_HEADERS


@pytest.mark.unit
def test_xlsx_gap_plan_renders_an_action_rows_owner_and_deadline() -> None:
    from openpyxl import load_workbook

    ctx = _ctx_with_actions()
    code = ctx.gap.gaps[0].code
    ctx = _ctx_with_actions(
        {
            code: _action(
                code,
                characterization="mitigate",
                owner="Priya Raman, CISO",
                deadline="2026-11-30",
                resources="2 FTE, $40k tooling",
                success_criteria="Policy approved and evidenced org-wide",
                poam_ref="POAM-2026-014",
            )
        }
    )
    ws = load_workbook(io.BytesIO(render_xlsx(ctx)))["Gap Plan"]
    row = [ws.cell(row=2, column=c).value for c in range(1, ws.max_column + 1)]
    assert row[0] == code
    cells = dict(zip(GAP_PLAN_HEADERS, row, strict=True))
    assert cells["Owner"] == "Priya Raman, CISO"
    assert cells["Deadline"] == "2026-11-30"
    assert cells["Characterization"] == "mitigate"
    assert cells["Resources"] == "2 FTE, $40k tooling"
    assert cells["Success criteria"] == "Policy approved and evidenced org-wide"
    assert cells["POA&M ref"] == "POAM-2026-014"


@pytest.mark.unit
def test_xlsx_gap_plan_priority_override_wins_over_the_computed_score() -> None:
    from openpyxl import load_workbook

    base = _ctx_with_actions()
    top = base.gap.gaps[0]
    computed = f"{top.priority_score:.2f}"
    # The assertion is only meaningful if the two values differ.
    assert computed != "P1", "fixture is vacuous: computed priority already reads P1"
    plain = load_workbook(io.BytesIO(render_xlsx(base)))["Gap Plan"]
    assert plain.cell(row=2, column=8).value == computed

    ctx = _ctx_with_actions({top.code: _action(top.code, priority_override="P1")})
    ws = load_workbook(io.BytesIO(render_xlsx(ctx)))["Gap Plan"]
    assert ws.cell(row=2, column=1).value == top.code
    assert ws.cell(row=2, column=8).value == "P1"


@pytest.mark.unit
def test_pdf_carries_the_tier_model_methodology_and_a_next_step_line() -> None:
    text = _pdf_norm(render_pdf(_ctx_with_actions()))
    partial = TIER_DEFINITIONS[0]
    adaptive = TIER_DEFINITIONS[-1]
    assert f"Tier 1 — {partial.short_label}" in text
    assert " ".join(partial.description.split()) in text
    assert " ".join(adaptive.description.split()) in text
    # Computed in Python from the gap rows: 20 shown gaps, all 1 tier short.
    assert (
        "Start with the 20 subcategory gap(s) sitting 1 tier(s) below target "
        "T4 (Adaptive) — they carry the largest lift." in text
    )
    assert "0 of 20 gap(s) in this action plan name an owner; assign the remaining 20." in text


@pytest.mark.unit
def test_docx_carries_the_action_plan_the_tier_model_and_a_next_step_line() -> None:
    ctx = _ctx_with_actions()
    code = ctx.gap.gaps[0].code
    ctx = _ctx_with_actions({code: _action(code, owner="Priya Raman, CISO", deadline="2026-11-30")})
    raw = render_docx(ctx)
    assert raw[:2] == b"PK"
    text = _docx_norm(raw)
    assert "Action plan" in text
    assert f"Tier 4 — {TIER_DEFINITIONS[-1].short_label}" in text
    assert " ".join(TIER_DEFINITIONS[-1].description.split()) in text
    assert (
        "Start with the 20 subcategory gap(s) sitting 1 tier(s) below target "
        "T4 (Adaptive) — they carry the largest lift." in text
    )
    assert "1 of 20 gap(s) in this action plan name an owner; assign the remaining 19." in text
    assert "Priya Raman, CISO" in text
    assert "2026-11-30" in text


@pytest.mark.unit
def test_next_steps_report_the_zero_gap_state() -> None:
    a, answers = _build_inputs(answers_tier=4)
    tier_map = {ans.subcategory_code: ans.maturity_tier for ans in answers}
    ctx = build_context(
        client_legal_name="Atlas Defense Solutions",
        service_title="NIST CSF 2.0 Assessment",
        assessment=a,
        answers=answers,
        score=compute_score(tier_map),
        gap=analyze_gaps(tier_map, target_tier=3),
    )
    expected = (
        "No subcategory scored below target T3 (Repeatable) — maintain the current "
        "controls and re-assess on the next cycle."
    )
    assert expected in _pdf_norm(render_pdf(ctx))
    assert expected in _docx_norm(render_docx(ctx))


FULL_COVERAGE_REASSURANCE = "maintain the current controls and re-assess on the next cycle"


def _partially_scored_ctx(scored: int, tier: int | None):
    """`scored` subcategories carry `tier`; the rest are genuinely unscored."""
    a, answers = _build_inputs(answers_tier=None)
    for ans in answers[:scored]:
        ans.maturity_tier = tier
    tier_map = {ans.subcategory_code: ans.maturity_tier for ans in answers}
    return build_context(
        client_legal_name="Atlas Defense Solutions",
        service_title="NIST CSF 2.0 Assessment",
        assessment=a,
        answers=answers,
        score=compute_score(tier_map),
        gap=analyze_gaps(tier_map, target_tier=3),
    )


@pytest.mark.unit
def test_next_steps_record_no_finding_when_nothing_is_scored() -> None:
    """Zero scored answers produce zero gaps, and a zero-gap all-clear over an
    unassessed framework is a finding of adequacy the data cannot support."""
    ctx = _partially_scored_ctx(0, None)
    assert ctx.score.answered_subcategories == 0
    assert ctx.gap.total_gap_count == 0
    for text in (_pdf_norm(render_pdf(ctx)), _docx_norm(render_docx(ctx))):
        assert "No subcategory has been scored, so this report records no maturity finding." in text
        assert f"All {len(SUBCATEGORIES)} subcategories in the NIST CSF 2.0 catalog" in text
        assert "remain unassessed." in text
        assert FULL_COVERAGE_REASSURANCE not in text


@pytest.mark.unit
def test_next_steps_name_the_unscored_count_when_coverage_is_partial() -> None:
    """3 of 106 scored, all at target: no scored gap, but 103 subcategories
    carry no finding and the report must not read as an all-clear."""
    ctx = _partially_scored_ctx(3, 3)
    unscored = len(SUBCATEGORIES) - 3
    assert ctx.score.answered_subcategories == 3
    assert ctx.gap.total_gap_count == 0
    for text in (_pdf_norm(render_pdf(ctx)), _docx_norm(render_docx(ctx))):
        assert "No scored subcategory fell below target T3 (Repeatable)." in text
        assert (
            f"{unscored} of {len(SUBCATEGORIES)} subcategories are unscored and carry "
            f"no finding." in text
        )
        assert FULL_COVERAGE_REASSURANCE not in text


@pytest.mark.unit
def test_xlsx_answers_tier_cells_carry_the_graded_fill() -> None:
    from openpyxl import load_workbook

    from app import export_style

    ws = load_workbook(io.BytesIO(render_xlsx(_ctx_with_actions())))["Answers"]
    want = "FF" + export_style.graded_hex(3, 4).lstrip("#").upper()
    cell = ws.cell(row=2, column=6)
    assert cell.value == 3
    assert cell.fill.start_color.rgb == want
    # A different tier must produce a different fill, or the ramp is inert.
    assert want != "FF" + export_style.graded_hex(1, 4).lstrip("#").upper()


@pytest.mark.unit
def test_xlsx_unscored_answer_rows_carry_no_fill() -> None:
    from openpyxl import load_workbook

    a, answers = _build_inputs(answers_tier=None)
    tier_map = {ans.subcategory_code: None for ans in answers}
    ctx = build_context(
        client_legal_name=None,
        service_title="x",
        assessment=a,
        answers=answers,
        score=compute_score(tier_map),
        gap=analyze_gaps(tier_map),
    )
    ws = load_workbook(io.BytesIO(render_xlsx(ctx)))["Answers"]
    assert ws.cell(row=2, column=7).value == "Unscored"
    assert ws.cell(row=2, column=6).fill.fill_type is None
    # The row still reaches the Evidence column, so "no tier" is not "no row".
    assert ws.cell(row=2, column=9).value == "No evidence attached"


@pytest.mark.unit
def test_xlsx_per_function_average_tier_cells_carry_the_graded_fill() -> None:
    from openpyxl import load_workbook

    from app import export_style

    ws = load_workbook(io.BytesIO(render_xlsx(_ctx_with_actions())))["Score Summary"]
    want = "FF" + export_style.graded_hex(3, 4).lstrip("#").upper()
    # Rows 1-6 metadata, 7 blank, 8 header, 9+ the six functions.
    assert ws.cell(row=8, column=6).value == "Average tier"
    for r in range(9, 15):
        cell = ws.cell(row=r, column=6)
        assert cell.value == "3.00", f"row {r} is not the all-tier-3 average"
        assert cell.fill.start_color.rgb == want, f"row {r} carries no graded fill"


@pytest.mark.unit
def test_xlsx_answers_evidence_renders_the_attached_filename() -> None:
    from openpyxl import load_workbook

    artifact_id = uuid.uuid4()
    a, answers = _build_inputs(answers_tier=3)
    answers[0].evidence_artifact_id = artifact_id
    tier_map = {ans.subcategory_code: ans.maturity_tier for ans in answers}
    ctx = build_context(
        client_legal_name="Atlas Defense Solutions",
        service_title="x",
        assessment=a,
        answers=answers,
        score=compute_score(tier_map),
        gap=analyze_gaps(tier_map, target_tier=4),
        evidence_names={artifact_id: "gv-oc-01-charter-2026.pdf"},
    )
    ws = load_workbook(io.BytesIO(render_xlsx(ctx)))["Answers"]
    assert ws.cell(row=1, column=9).value == "Evidence"
    row = {ws.cell(row=r, column=1).value: r for r in range(2, ws.max_row + 1)}
    assert ws.cell(row=row[answers[0].subcategory_code], column=9).value == (
        "gv-oc-01-charter-2026.pdf"
    )


@pytest.mark.unit
def test_xlsx_answers_evidence_renders_the_null_sentence() -> None:
    from openpyxl import load_workbook

    ws = load_workbook(io.BytesIO(render_xlsx(_ctx_with_actions())))["Answers"]
    assert ws.cell(row=2, column=9).value == "No evidence attached"


@pytest.mark.unit
def test_build_context_refuses_an_unresolved_evidence_pointer() -> None:
    """FAIL LOUDLY: a dangling pointer must never degrade into the NULL sentence."""
    a, answers = _build_inputs(answers_tier=3)
    answers[0].evidence_artifact_id = uuid.uuid4()
    tier_map = {ans.subcategory_code: ans.maturity_tier for ans in answers}
    with pytest.raises(ValueError, match="evidence_names is missing"):
        build_context(
            client_legal_name="Atlas Defense Solutions",
            service_title="x",
            assessment=a,
            answers=answers,
            score=compute_score(tier_map),
            gap=analyze_gaps(tier_map, target_tier=4),
            evidence_names={},
        )


# ---------------------------------------------------------------------------
# S9: PDF acceptance contract — section order plus one score/model/action link
# ---------------------------------------------------------------------------


def _assert_section_sequence(text: str, sequence: list[str]) -> None:
    """Assert every string appears in the rendered text, in exactly this order.

    Each needle is searched for only AFTER the previous match, so this is a true
    subsequence check rather than a set of `in` checks. A section that renders
    but lands in the wrong place still reads as a defensible report to `in`, and
    that is the failure this is here to catch.
    """
    last = -1
    last_needle = "(start of document)"
    for needle in sequence:
        found = text.find(needle, last + 1)
        assert found != -1, (
            f"{needle!r} does not appear in the rendered PDF after "
            f"{last_needle!r} (index {last})"
        )
        last, last_needle = found, needle


@pytest.mark.unit
def test_pdf_acceptance_contract_orders_sections_and_links_the_tier_to_its_action() -> None:
    """The CSF PDF's acceptance contract, over real rendered bytes.

    Section order: Maturity summary -> How these tiers are scored -> Per-function
    rollup -> Top remediation gaps -> Action plan.

    Representative linkage, all on the SAME subcategory: the average tier (the
    score), the NIST tier definition that says what that tier means (its
    citation reference — the model the number is read against), the gap row
    showing its ramp to target, and the action-plan row naming an owner and a
    deadline (its action). The tier model has to precede the gap list: a reader
    cannot judge a T3 -> T4 gap before being told what T3 and T4 are.
    """
    from datetime import date

    def _tier_def(tier: int):
        for d in TIER_DEFINITIONS:
            if int(d.tier) == tier:
                return d
        raise AssertionError(f"no tier definition for T{tier}")

    action = _action(
        "DE.AE-02",
        owner="Dana Iyer",
        deadline=date(2026, 11, 30),
        characterization="Ongoing",
    )
    ctx = _ctx_with_actions({"DE.AE-02": action})
    text = _pdf_norm(render_pdf(ctx))
    first_gap = ctx.gap.gaps[0]
    assert first_gap.code == "DE.AE-02", "the action must key the first shown gap"

    _assert_section_sequence(
        text,
        [
            # Section 1, and the score.
            "Maturity summary",
            "Overall maturity: Repeatable",
            "Average tier: 3.00",
            # Section 2: the model the score is read against, from TIER_DEFINITIONS.
            "How these tiers are scored",
            _tier_def(3).description,
            _tier_def(4).description,
            # Section 3.
            "Per-function rollup",
            # Section 4, and the gap's ramp on the same subcategory.
            f"Top remediation gaps (target T{ctx.gap.target_tier})",
            f"{first_gap.code} {first_gap.function.value}",
            f"T{first_gap.current_tier} → T{first_gap.target_tier}",
            # Section 5, and the action that closes it, with its owner and date.
            f"Action plan ({len(ctx.gap.gaps)} of {ctx.gap.total_gap_count} gaps shown)",
            "Dana Iyer",
            "2026-11-30",
        ],
    )
