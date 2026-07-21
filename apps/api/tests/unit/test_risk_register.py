"""Risk Register: gate, generate, tier-from-code, link validation (Work Order E)."""

from __future__ import annotations

import io
import os
from collections.abc import Iterator
from pathlib import Path
from types import SimpleNamespace

import pytest
from alembic import command
from alembic.config import Config
from fastapi.testclient import TestClient
from sqlalchemy import create_engine
from sqlalchemy.orm import Session, sessionmaker

from app.ai.llm import FixtureProvider, LLMClient, LLMResponse
from app.risk.exporters import build_context as build_risk_context
from app.risk.exporters import render_docx, render_pdf


@pytest.fixture()
def app_client(tmp_path) -> Iterator[tuple[TestClient, FixtureProvider]]:
    url = f"sqlite:///{tmp_path / 'shield-risk.db'}"
    os.environ["DATABASE_URL"] = url
    api_root = Path(__file__).resolve().parents[2]
    cfg = Config(str(api_root / "alembic.ini"))
    cfg.set_main_option("script_location", str(api_root / "alembic"))
    cfg.set_main_option("sqlalchemy.url", url)
    command.upgrade(cfg, "head")
    engine = create_engine(url, future=True)
    TestSession = sessionmaker(bind=engine, autoflush=False, autocommit=False, future=True)

    from app.db.session import get_db
    from app.main import create_app
    from app.routes.artifacts import _storage_dep
    from app.routes.risk import _llm_dep
    from app.storage.local import LocalFilesystemStorage

    def override_get_db() -> Iterator[Session]:
        db = TestSession()
        try:
            yield db
        finally:
            db.close()

    provider = FixtureProvider()
    storage = LocalFilesystemStorage(tmp_path / "storage")
    app = create_app()
    app.dependency_overrides[get_db] = override_get_db
    app.dependency_overrides[_llm_dep] = lambda: LLMClient(provider)
    app.dependency_overrides[_storage_dep] = lambda: storage
    with TestClient(app) as c:
        yield c, provider


def _admin(c: TestClient) -> tuple[str, str]:
    admin = c.post(
        "/auth/register",
        json={
            "email": "admin@kentro.example",
            "password": "correct horse battery staple!",
            "display_name": "A",
        },
    )
    bearer = admin.json()["tokens"]["access_token"]
    cid = c.post(
        "/admin/clients",
        headers={"Authorization": f"Bearer {bearer}"},
        json={"legal_name": "Acme"},
    ).json()["id"]
    return bearer, cid


def _seed_attack_and_zt(c: TestClient, bearer: str, cid: str) -> tuple[str, str]:
    """Returns (a gap technique_code, a ZT capability_code)."""
    h = {"Authorization": f"Bearer {bearer}", "X-Client-Id": cid}
    asvc = c.post(
        "/attack/services", headers=h, json={"kind": "attack_coverage", "title": "ATT&CK"}
    )
    a = c.post(f"/attack/services/{asvc.json()['id']}/assessments", headers=h)
    cov = a.json()["coverage"][0]
    technique = cov["technique_code"]
    c.patch(f"/attack/coverage/{cov['id']}", headers=h, json={"status": "gap"})

    zsvc = c.post("/zt/services", headers=h, json={"kind": "zero_trust_cisa", "title": "ZT"})
    za = c.post(f"/zt/services/{zsvc.json()['id']}/assessments", headers=h)
    zans = za.json()["answers"][0]
    capability = zans["capability_code"]
    c.patch(f"/zt/answers/{zans['id']}", headers=h, json={"maturity_stage": 1})
    return technique, capability


@pytest.mark.unit
def test_gate_locked_without_attack(app_client) -> None:
    c, _ = app_client
    bearer, cid = _admin(c)
    bh = {"Authorization": f"Bearer {bearer}"}
    g = c.get(f"/risk/clients/{cid}/gate", headers=bh)
    assert g.status_code == 200
    assert g.json()["unlocked"] is False
    # Generate refuses while locked.
    r = c.post(f"/risk/clients/{cid}/register/generate", headers=bh)
    assert r.status_code == 409


@pytest.mark.unit
def test_generate_derives_tier_in_code_and_validates_links(app_client) -> None:
    c, provider = app_client
    bearer, cid = _admin(c)
    technique, capability = _seed_attack_and_zt(c, bearer, cid)
    bh = {"Authorization": f"Bearer {bearer}"}

    assert c.get(f"/risk/clients/{cid}/gate", headers=bh).json()["unlocked"] is True

    provider.register_static(
        "risk_synthesize",
        LLMResponse(
            '{"entries": [{"title": "Credential theft exposure",'
            ' "description": "EDR gap", "axis": "detection",'
            ' "source": "coverage_finding", "source_id": "' + technique + '",'
            ' "linked_techniques": ["' + technique + '", "T9999"],'
            ' "linked_controls": ["' + capability + '", "BOGUS.XX.01"],'
            ' "likelihood": "high", "impact": "catastrophic",'
            ' "recommended_action": "remediate", "rationale": "...",'
            ' "tier": "low"}]}'  # AI's "tier" must be ignored.
        ),
    )

    r = c.post(f"/risk/clients/{cid}/register/generate", headers=bh)
    assert r.status_code == 201, r.text
    body = r.json()
    assert body["version"] == 1
    assert len(body["entries"]) == 1
    e = body["entries"][0]
    # Tier is code-derived (High + Catastrophic -> Critical), NOT the AI's "low".
    assert e["tier"] == "critical"
    # Invented technique/control dropped; only the real ones remain.
    assert e["linked_techniques"] == [technique]
    assert e["linked_controls"] == [capability]
    assert e["origin"] == "ai_generated"
    assert body["tier_counts"]["critical"] == 1
    assert body["axis_counts"]["detection"] == 1


@pytest.mark.unit
def test_export_renders_and_stores_three_files(app_client) -> None:
    c, provider = app_client
    bearer, cid = _admin(c)
    technique, capability = _seed_attack_and_zt(c, bearer, cid)
    bh = {"Authorization": f"Bearer {bearer}"}
    provider.register_static(
        "risk_synthesize",
        LLMResponse(
            '{"entries": [{"title": "Risk one", "axis": "detection",'
            ' "likelihood": "high", "impact": "catastrophic",'
            ' "recommended_action": "remediate"}]}'
        ),
    )
    c.post(f"/risk/clients/{cid}/register/generate", headers=bh)
    r = c.post(f"/risk/clients/{cid}/register/export", headers=bh)
    assert r.status_code == 200, r.text
    body = r.json()
    assert body["xlsx_filename"].endswith(".xlsx")
    assert body["pdf_filename"].endswith(".pdf")
    assert body["docx_filename"].endswith(".docx")
    # Each downloads as a real file (artifact download is tenant-scoped, so the
    # admin names the active client via X-Client-Id).
    dh = {**bh, "X-Client-Id": cid}
    xlsx = c.get(f"/artifacts/{body['xlsx_artifact_id']}/download", headers=dh)
    assert xlsx.status_code == 200 and xlsx.content[:2] == b"PK"
    pdf = c.get(f"/artifacts/{body['pdf_artifact_id']}/download", headers=dh)
    assert pdf.status_code == 200 and pdf.content.startswith(b"%PDF-")
    docx = c.get(f"/artifacts/{body['docx_artifact_id']}/download", headers=dh)
    assert docx.status_code == 200 and docx.content[:2] == b"PK"


# ---------------------------------------------------------------------------
# Exporter content (pure renderers — no DB). The XLSX is already content-tested
# via the download in test_export_renders_and_stores_three_files; here we prove
# the PDF + DOCX carry the title and a known entry (SMOKE §10).
# ---------------------------------------------------------------------------


def _risk_entry() -> SimpleNamespace:
    return SimpleNamespace(
        title="Credential theft exposure",
        description="EDR coverage gap",
        axis="detection",
        source="coverage_finding",
        source_id="T1078",
        linked_techniques=["T1078"],
        linked_controls=["ID.PA.01"],
        likelihood="high",
        impact="catastrophic",
        tier="critical",
        compensating_controls=None,
        residual_risk=None,
        recommended_action="remediate",
        rationale="Primary detection layer absent.",
        origin="ai_generated",
        trust=None,
    )


def _pdf_text(raw: bytes) -> str:
    from pypdf import PdfReader

    reader = PdfReader(io.BytesIO(raw))
    return "".join(page.extract_text() for page in reader.pages)


@pytest.mark.unit
def test_pdf_carries_title_client_and_a_known_entry() -> None:
    ctx = build_risk_context(
        client_legal_name="Atlas Defense Solutions", version=3, entries=[_risk_entry()]
    )
    raw = render_pdf(ctx)
    assert raw.startswith(b"%PDF-")
    text = _pdf_text(raw)
    assert "Risk Register" in text  # document title
    assert "Atlas Defense Solutions" in text  # client name
    assert "Credential theft exposure" in text  # a known register entry


@pytest.mark.unit
def test_docx_carries_title_client_and_a_known_entry() -> None:
    from docx import Document

    ctx = build_risk_context(
        client_legal_name="Atlas Defense Solutions", version=3, entries=[_risk_entry()]
    )
    raw = render_docx(ctx)
    assert raw[:2] == b"PK"  # docx is a zip envelope
    doc = Document(io.BytesIO(raw))
    paras = [p.text for p in doc.paragraphs if p.text]
    assert "Risk Register (v3)" in paras  # title heading
    assert "Atlas Defense Solutions" in paras  # client subtitle
    cells = [c.text for t in doc.tables for row in t.rows for c in row.cells]
    assert "Credential theft exposure" in cells  # a known register entry


@pytest.mark.unit
def test_each_generate_is_a_new_version(app_client) -> None:
    c, provider = app_client
    bearer, cid = _admin(c)
    _seed_attack_and_zt(c, bearer, cid)
    bh = {"Authorization": f"Bearer {bearer}"}
    provider.register_static("risk_synthesize", LLMResponse('{"entries": []}'))

    v1 = c.post(f"/risk/clients/{cid}/register/generate", headers=bh).json()
    v2 = c.post(f"/risk/clients/{cid}/register/generate", headers=bh).json()
    assert v1["version"] == 1
    assert v2["version"] == 2
    latest = c.get(f"/risk/clients/{cid}/register/latest", headers=bh)
    assert latest.json()["version"] == 2
