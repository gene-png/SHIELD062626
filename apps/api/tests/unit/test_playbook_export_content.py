"""Content assertions over the five CSF full-Playbook renderers.

The renderers in `app.csf.playbook_export` are pure functions over
`enterprise_rows` / `tier_profiles` / `gap_actions` — no DB, no TestClient. We
construct lightweight rows (attribute-access only, so `SimpleNamespace` is
enough) and read the rendered bytes back:

  - XLSX  → openpyxl (the SMOKE §19 Action Plan contract).
  - PDF   → pypdf.PdfReader.extract_text() (distinctive substrings only —
            reportlab text extraction is whitespace-mangled, so we never assert
            exact layout).
  - DOCX  → docx.Document over BytesIO (heading + table cell text; colours /
            shading are NOT assertable, values only).

These lock the content the §10 human eyeball used to certify by hand.
"""

from __future__ import annotations

import io
from types import SimpleNamespace

import pytest

from app.csf.playbook import gap_priority
from app.csf.playbook_export import (
    render_exec_docx,
    render_exec_pdf,
    render_full_docx,
    render_full_pdf,
    render_xlsx,
)

CLIENT = "Atlas Defense Solutions"
DOC_TITLE = "NIST CSF 2.0"

# The three functions we exercise, each mapped to its display name so the
# scorecard/function-detail text is predictable.
GOVERN_DEFAULT_PRIORITY = gap_priority(is_core=True, high_tier=True, multi_system=True)  # -> P1
ASSET_DEFAULT_PRIORITY = gap_priority(is_core=False, high_tier=False, multi_system=False)  # -> P3


def _enterprise_rows() -> list[SimpleNamespace]:
    """Three enterprise subcategories across GV / ID / PR — two gaps, one met."""
    return [
        SimpleNamespace(
            subcategory_code="GV.OC-01",
            name="Organizational context",
            function="GV",
            tier_levels={"high": 2, "moderate": 3, "low": 4},
            enterprise_level=2,
            rollup_rule=6,
            target_level=4,
            gap=True,
            priority=GOVERN_DEFAULT_PRIORITY,
        ),
        SimpleNamespace(
            subcategory_code="ID.AM-02",
            name="Asset inventory",
            function="ID",
            tier_levels={"high": 1, "moderate": 2, "low": 2},
            enterprise_level=1,
            rollup_rule=3,
            target_level=3,
            gap=True,
            priority=ASSET_DEFAULT_PRIORITY,
        ),
        SimpleNamespace(
            subcategory_code="PR.AA-01",
            name="Identity management",
            function="PR",
            tier_levels={"high": 5, "moderate": 5, "low": 5},
            enterprise_level=5,
            rollup_rule=1,
            target_level=4,
            gap=False,
            priority=None,
        ),
    ]


def _tier_rows() -> list[SimpleNamespace]:
    return [
        SimpleNamespace(
            subcategory_code="GV.OC-01",
            governance=2,
            policy=1,
            implementation=1,
            monitoring=1,
            improvement=0,
            total=5,
            level=2,
            evidence_capped=True,
            in_scope=True,
            target_level=4,
        )
    ]


def _pdf_text(raw: bytes) -> str:
    from pypdf import PdfReader

    reader = PdfReader(io.BytesIO(raw))
    return "".join(page.extract_text() for page in reader.pages)


def _docx_paragraphs(raw: bytes) -> list[str]:
    from docx import Document

    return [p.text for p in Document(io.BytesIO(raw)).paragraphs if p.text]


def _docx_table_cells(raw: bytes) -> list[list[str]]:
    from docx import Document

    doc = Document(io.BytesIO(raw))
    return [[c.text for row in t.rows for c in row.cells] for t in doc.tables]


# ---------------------------------------------------------------------------
# XLSX — the SMOKE §19 Action Plan contract
# ---------------------------------------------------------------------------

ACTION_PLAN_HEADERS = [
    "Subcategory",
    "Outcome",
    "Enterprise",
    "Target",
    "Priority",
    "Characterization",
    "Owner",
    "Deadline",
    "Resources",
    "Success criteria",
    "POA&M ref",
]


@pytest.mark.unit
def test_xlsx_action_plan_sheet_has_expected_headers() -> None:
    from openpyxl import load_workbook

    raw = render_xlsx(
        client_name=CLIENT,
        version=7,
        enterprise_rows=_enterprise_rows(),
        tier_profiles={"high": _tier_rows()},
    )
    wb = load_workbook(io.BytesIO(raw))
    assert "Action Plan" in wb.sheetnames
    ws = wb["Action Plan"]
    header = [ws.cell(row=1, column=c).value for c in range(1, len(ACTION_PLAN_HEADERS) + 1)]
    assert header == ACTION_PLAN_HEADERS


@pytest.mark.unit
def test_xlsx_action_plan_only_lists_gaps() -> None:
    from openpyxl import load_workbook

    wb = load_workbook(
        io.BytesIO(
            render_xlsx(
                client_name=CLIENT,
                version=7,
                enterprise_rows=_enterprise_rows(),
                tier_profiles={"high": _tier_rows()},
            )
        )
    )
    ws = wb["Action Plan"]
    codes = [ws.cell(row=r, column=1).value for r in range(2, ws.max_row + 1)]
    # Two gaps present; the met subcategory (PR.AA-01) is excluded.
    assert codes == ["GV.OC-01", "ID.AM-02"]


@pytest.mark.unit
def test_xlsx_action_plan_priority_defaults_from_gap_priority() -> None:
    """No stored override -> the Priority cell is the code-computed default."""
    from openpyxl import load_workbook

    wb = load_workbook(
        io.BytesIO(
            render_xlsx(
                client_name=CLIENT,
                version=7,
                enterprise_rows=_enterprise_rows(),
                tier_profiles={"high": _tier_rows()},
                gap_actions=None,
            )
        )
    )
    ws = wb["Action Plan"]
    by_code = {ws.cell(row=r, column=1).value: ws.cell(row=r, column=5).value for r in (2, 3)}
    assert by_code["GV.OC-01"] == GOVERN_DEFAULT_PRIORITY == "P1"
    assert by_code["ID.AM-02"] == ASSET_DEFAULT_PRIORITY == "P3"


@pytest.mark.unit
def test_xlsx_action_plan_priority_override_wins() -> None:
    """A stored priority_override beats the code-computed default."""
    from openpyxl import load_workbook

    actions = {
        "ID.AM-02": SimpleNamespace(
            priority_override="P1",
            characterization="mitigate",
            owner="Jane Ops",
            deadline="2026-09-30",
            resources="2 FTE",
            success_criteria="inventory at 95%",
            poam_ref="POAM-42",
        )
    }
    wb = load_workbook(
        io.BytesIO(
            render_xlsx(
                client_name=CLIENT,
                version=7,
                enterprise_rows=_enterprise_rows(),
                tier_profiles={"high": _tier_rows()},
                gap_actions=actions,
            )
        )
    )
    ws = wb["Action Plan"]
    row = {h: ws.cell(row=3, column=i + 1).value for i, h in enumerate(ACTION_PLAN_HEADERS)}
    assert row["Subcategory"] == "ID.AM-02"
    # Override "P1" wins over the code default "P3".
    assert ASSET_DEFAULT_PRIORITY == "P3"
    assert row["Priority"] == "P1"
    # The stored annotation columns are carried through too.
    assert row["Characterization"] == "mitigate"
    assert row["Owner"] == "Jane Ops"
    assert row["POA&M ref"] == "POAM-42"


# ---------------------------------------------------------------------------
# PDF — executive briefing + full playbook
# ---------------------------------------------------------------------------


@pytest.mark.unit
def test_exec_pdf_carries_client_title_function_and_gap() -> None:
    raw = render_exec_pdf(client_name=CLIENT, version=7, enterprise_rows=_enterprise_rows())
    assert raw.startswith(b"%PDF-")
    text = _pdf_text(raw)
    assert CLIENT in text  # client name
    assert DOC_TITLE in text  # document title
    assert "Govern" in text  # at least one CSF function name
    # A known gap value — the outcome + its code + its computed priority.
    assert "GV.OC-01" in text
    assert "Organizational context" in text


@pytest.mark.unit
def test_full_pdf_carries_client_title_function_and_gap() -> None:
    raw = render_full_pdf(client_name=CLIENT, version=7, enterprise_rows=_enterprise_rows())
    assert raw.startswith(b"%PDF-")
    text = _pdf_text(raw)
    assert CLIENT in text
    assert DOC_TITLE in text
    # Full playbook adds the methodology + per-function detail.
    assert "Methodology" in text
    assert "Identify" in text
    assert "Asset inventory" in text


# ---------------------------------------------------------------------------
# DOCX — executive briefing + full playbook (values only, never shading)
# ---------------------------------------------------------------------------

SCORECARD_HEADERS = {"Function", "Subcategories", "Maturity", "Target", "Gaps"}


@pytest.mark.unit
def test_exec_docx_heading_scorecard_and_maturity_cell() -> None:
    raw = render_exec_docx(client_name=CLIENT, version=7, enterprise_rows=_enterprise_rows())
    paras = _docx_paragraphs(raw)
    assert DOC_TITLE in paras
    assert "Executive summary" in paras
    assert f"Prepared for: {CLIENT}" in paras

    scorecard = next(t for t in _docx_table_cells(raw) if set(t) >= SCORECARD_HEADERS)
    # A known maturity cell value: Govern rolled up to L2.
    assert "Govern" in scorecard
    assert "L2" in scorecard


@pytest.mark.unit
def test_full_docx_heading_scorecard_and_maturity_cell() -> None:
    raw = render_full_docx(client_name=CLIENT, version=7, enterprise_rows=_enterprise_rows())
    paras = _docx_paragraphs(raw)
    assert DOC_TITLE in paras
    assert "2. Methodology" in paras
    assert "Govern (GV)" in paras

    scorecard = next(t for t in _docx_table_cells(raw) if set(t) >= SCORECARD_HEADERS)
    assert "Protect" in scorecard
    # Protect meets its target, rolled up to L5.
    assert "L5" in scorecard


# ---------------------------------------------------------------------------
# Target coverage: "no target recorded" is not "target met" (H1)
#
# `routes/csf.py` resolves `gap` to False whenever a subcategory has no target
# level, so an assessment with no targets set produces an empty gap list that is
# indistinguishable from one where every target is met. These five renderers
# then told the client every subcategory met its target. `target_level` is
# nullable and explicitly written as None, so the distinction is a fact the
# exporter can read; these tests pin that it does.
#
# The `csf/exporters.py` deliverable already discriminates on the ANSWERED axis
# (`_no_gap_steps`, S3). This is the same discipline on the TARGET axis, which
# is the axis this file's gap flag actually depends on.
# ---------------------------------------------------------------------------

# The claim that must never appear on a report with no targets recorded.
FALSE_ALL_CLEAR = "every in-scope subcategory meets its target"
# The advice that must never follow it.
FALSE_ADVICE = "Maintain current controls"


def _flat(text: str) -> str:
    """Collapse the whitespace reportlab's text extraction mangles, so a phrase
    assertion tests the words rather than the line breaking."""
    return " ".join(text.split())


def _rows_without_targets() -> list[SimpleNamespace]:
    """The default state of a real engagement: every subcategory scored, no
    target level recorded anywhere.

    This is not a contrived empty input. `models/csf_profile.py` types
    `target_level` as nullable and `routes/csf.py` writes None explicitly, so an
    engagement that has been fully scored but never had targets set lands here.
    """
    return [
        SimpleNamespace(
            subcategory_code=row.subcategory_code,
            name=row.name,
            function=row.function,
            tier_levels=row.tier_levels,
            enterprise_level=row.enterprise_level,
            rollup_rule=row.rollup_rule,
            target_level=None,
            gap=False,
            priority=None,
        )
        for row in _enterprise_rows()
    ]


def _rows_with_partial_targets() -> list[SimpleNamespace]:
    """One of three subcategories carries a target, and it is met."""
    rows = _rows_without_targets()
    rows[2].target_level = 4
    return rows


def _rows_all_targets_met() -> list[SimpleNamespace]:
    """Every subcategory carries a target and every target is met. This is the
    only input for which the reassuring line is true."""
    rows = _enterprise_rows()
    for row in rows:
        row.target_level = 1
        row.gap = False
        row.priority = None
    return rows


@pytest.mark.unit
@pytest.mark.parametrize("render", [render_exec_pdf, render_full_pdf])
def test_pdf_with_no_targets_recorded_does_not_claim_targets_are_met(render) -> None:
    text = _flat(
        _pdf_text(render(client_name=CLIENT, version=7, enterprise_rows=_rows_without_targets()))
    )
    assert FALSE_ALL_CLEAR not in text
    assert FALSE_ADVICE not in text
    # And it says what is actually true instead.
    assert "no target maturity level has been recorded" in text.lower()


@pytest.mark.unit
@pytest.mark.parametrize("render", [render_exec_docx, render_full_docx])
def test_docx_with_no_targets_recorded_does_not_claim_targets_are_met(render) -> None:
    paras = _docx_paragraphs(
        render(client_name=CLIENT, version=7, enterprise_rows=_rows_without_targets())
    )
    joined = _flat(" ".join(paras))
    assert FALSE_ALL_CLEAR not in joined
    assert FALSE_ADVICE not in joined
    assert "no target maturity level has been recorded" in joined.lower()


@pytest.mark.unit
@pytest.mark.parametrize("render", [render_exec_pdf, render_full_pdf])
def test_pdf_with_no_targets_recorded_does_not_report_zero_shortfalls(render) -> None:
    """`0 subcategories fall short of their target maturity` reads as an
    all-clear when the truth is that nothing could be measured."""
    text = _flat(
        _pdf_text(render(client_name=CLIENT, version=7, enterprise_rows=_rows_without_targets()))
    )
    assert "0 subcategories fall short" not in text


@pytest.mark.unit
def test_partial_target_coverage_names_the_uncovered_count() -> None:
    """One target of three, and it is met. The report may not generalise from
    the one to the three."""
    text = _flat(
        _pdf_text(
            render_full_pdf(
                client_name=CLIENT, version=7, enterprise_rows=_rows_with_partial_targets()
            )
        )
    )
    assert FALSE_ALL_CLEAR not in text
    assert FALSE_ADVICE not in text
    # 2 of the 3 carry no target and therefore carry no finding.
    assert "2 of 3" in text


@pytest.mark.unit
def test_full_target_coverage_still_gives_the_reassuring_line() -> None:
    """The counterexample that keeps the fix from being a blanket deletion: when
    every subcategory really does have a target and really does meet it, the
    reassuring sentence is true and must survive.

    Uses the executive renderer deliberately. `render_full_pdf` has no
    next-steps section, so it can never carry this advice and asserting on it
    there would pass for the wrong reason.
    """
    text = _flat(
        _pdf_text(
            render_exec_pdf(client_name=CLIENT, version=7, enterprise_rows=_rows_all_targets_met())
        )
    )
    assert FALSE_ADVICE in text
    assert "no target maturity level has been recorded" not in text.lower()
